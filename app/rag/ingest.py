from __future__ import annotations

import io
import logging
import os
import re
from functools import lru_cache
        codex/refactor-modules-to-remove-codex-markers
from typing import Dict, List, Optional, Protocol

from docx import Document
from pypdf import PdfReader

try:  # pragma: no cover - dependency provided in production image
    import tiktoken
except ImportError:  # pragma: no cover - fallback used in tests
    tiktoken = None  # type: ignore[assignment]


LOGGER = logging.getLogger(__name__)

from typing import Dict, Iterable, List, Optional, Protocol

try:  # pragma: no cover - dependency provided via requirements
    import tiktoken
except ImportError:  # pragma: no cover - used in tests when dependency missing
    tiktoken = None  # type: ignore
        main

from docx import Document
from pypdf import PdfReader

LOGGER = logging.getLogger(__name__)


class _Tokenizer(Protocol):
    """Subset of the tiktoken ``Encoding`` interface used by the tests."""

    def encode(self, text: str) -> List[int]:
        ...

    def decode(self, tokens: List[int]) -> str:
        ...


class _CharTokenizer:
        codex/refactor-modules-to-remove-codex-markers
    """Very small tokenizer used as a last-resort fallback."""

    """Fallback tokenizer that operates on individual characters."""
        main

    def encode(self, text: str) -> List[int]:
        return [ord(ch) for ch in text]

    def decode(self, tokens: List[int]) -> str:
        return "".join(chr(token) for token in tokens)


_TOKENIZER: Optional[_Tokenizer] = None


        codex/refactor-modules-to-remove-codex-markers
def _load_tiktoken(name: str) -> Optional[_Tokenizer]:
    if tiktoken is None:  # pragma: no cover - handled in tests
        return None
    try:
        return tiktoken.get_encoding(name)
    except Exception as exc:  # pragma: no cover - defensive fallback
        LOGGER.warning("Failed to load tokenizer '%s': %s", name, exc)
        try:
            return tiktoken.encoding_for_model(name)
        except Exception:  # pragma: no cover - final fallback
            return None


@lru_cache(maxsize=1)
def _default_tokenizer() -> _Tokenizer:
    name = os.getenv("RAG_TOKENIZER_NAME", "cl100k_base")
    tokenizer = _load_tiktoken(name)
    if tokenizer is None and name != "text-embedding-3-small":
        tokenizer = _load_tiktoken("text-embedding-3-small")
    return tokenizer or _CharTokenizer()

@lru_cache(maxsize=1)
def _byte_fallback() -> "tiktoken.core.Encoding":  # type: ignore[name-defined]
    """Return a byte-level tokenizer compatible with the tiktoken API."""

    mergeable_ranks = {bytes([i]): i for i in range(256)}
    return tiktoken.Encoding(  # type: ignore[call-arg]
        name="byte_fallback",
        pat_str=r"(?s:.)",
        mergeable_ranks=mergeable_ranks,
        special_tokens={},
    )
        main


def _get_tokenizer() -> _Tokenizer:
    """Return the tokenizer used to measure token lengths."""

    global _TOKENIZER
        codex/refactor-modules-to-remove-codex-markers
    if _TOKENIZER is None:
        _TOKENIZER = _default_tokenizer()
    return _TOKENIZER

    if _TOKENIZER is not None:
        return _TOKENIZER

    if tiktoken is None:
        _TOKENIZER = _CharTokenizer()
        return _TOKENIZER

    name = os.getenv("RAG_TOKENIZER_NAME", "cl100k_base")
    try:
        _TOKENIZER = tiktoken.get_encoding(name)
    except Exception:  # pragma: no cover - defensive fall-back
        try:
            _TOKENIZER = tiktoken.encoding_for_model("text-embedding-3-small")
        except Exception:  # pragma: no cover - fallback for unusual environments
            _TOKENIZER = _byte_fallback()
    return _TOKENIZER


def _clean(text: str) -> str:
    """Normalise whitespace extracted from source documents."""
        main

    return re.sub(r"\s+", " ", text).strip()

        codex/refactor-modules-to-remove-codex-markers
def _clean(text: str) -> str:
    """Collapse whitespace and trim the provided *text*."""

    return re.sub(r"\s+", " ", text).strip()


def _chunk(
    text: str,
    *,
    chunk: int = 900,
    overlap: int = 140,
    encoder: Optional[_Tokenizer] = None,
) -> List[str]:
    """Split *text* into token aware windows."""



def _normalise_window_size(value: int, minimum: int = 1) -> int:
    value = int(value)
    return minimum if value < minimum else value


def _normalise_overlap(chunk: int, overlap: int) -> int:
    overlap = 0 if overlap < 0 else int(overlap)
    if chunk <= 1:
        return 0
    return min(overlap, chunk - 1)


def _chunk(
    text: str,
    *,
    chunk: int = 900,
    overlap: int = 140,
    encoder: Optional[_Tokenizer] = None,
) -> List[str]:
    """Split ``text`` into overlapping windows based on token counts."""

        main
    if not text:
        return []

    tokenizer = encoder or _get_tokenizer()
        codex/refactor-modules-to-remove-codex-markers
    tokens = tokenizer.encode(text)
    if not tokens:
        return []

    chunk = max(int(chunk), 1)
    overlap = max(int(overlap), 0)
    if overlap >= chunk:
        overlap = chunk - 1 if chunk > 1 else 0

    token_ids = tokenizer.encode(text)
    if not token_ids:
        return []

    window = _normalise_window_size(chunk)
    step_overlap = _normalise_overlap(window, overlap)
        main

    pieces: List[str] = []
    start = 0
    total = len(tokens)
    while start < total:
        codex/refactor-modules-to-remove-codex-markers
        end = min(start + chunk, total)
        window_tokens = tokens[start:end]
        pieces.append(tokenizer.decode(window_tokens))

        end = min(start + window, total)
        tokens = token_ids[start:end]
        pieces.append(tokenizer.decode(tokens))
        main
        if end >= total:
            break
        next_start = end - step_overlap
        if next_start <= start:
            next_start = start + 1
        start = next_start

    return pieces


        codex/refactor-modules-to-remove-codex-markers
def _parse_pdf(data: bytes) -> List[tuple[int, str]]:
    reader = PdfReader(io.BytesIO(data))
    pages: List[tuple[int, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # pragma: no cover - defensive for exotic PDFs
            text = ""
        pages.append((index, _clean(text)))
    return pages


def _parse_docx(data: bytes) -> List[tuple[int, str]]:
    document = Document(io.BytesIO(data))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return [(1, _clean(text))]


def _iter_pdf_text(data: bytes) -> Iterable[tuple[int, str]]:
    reader = PdfReader(io.BytesIO(data))
    for page_number, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:  # pragma: no cover - pypdf quirks
            text = ""
        cleaned = _clean(text)
        if cleaned:
            yield page_number, cleaned


def _iter_docx_text(data: bytes) -> Iterable[tuple[int, str]]:
    document = Document(io.BytesIO(data))
    text = _clean("\n".join(paragraph.text for paragraph in document.paragraphs))
    if text:
        yield 1, text
        main

def _parse_txt(data: bytes) -> List[tuple[int, str]]:
    text = data.decode("utf-8", errors="ignore")
    return [(1, _clean(text))]

        codex/refactor-modules-to-remove-codex-markers

def _iter_txt_text(data: bytes) -> Iterable[tuple[int, str]]:
    text = _clean(data.decode("utf-8", errors="ignore"))
    if text:
        yield 1, text
        main

def parse_and_chunk(filename: str, data: bytes) -> List[Dict[str, object]]:
    """Parse *data* originating from *filename* and return chunk payloads."""

        codex/refactor-modules-to-remove-codex-markers
    name = (filename or "").strip()
    if not name:
        return []
    ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""

    if ext == "pdf":
        pages = _parse_pdf(data)
    elif ext == "docx":
        pages = _parse_docx(data)
    elif ext == "txt":
        pages = _parse_txt(data)
    else:
        return []

    chunk_size = int(os.getenv("RAG_CHUNK", "900"))
    overlap = int(os.getenv("RAG_OVERLAP", "140"))
    tokenizer = _get_tokenizer()

    chunks: List[Dict[str, object]] = []
    for page_number, page_text in pages:
        if not page_text:
            continue
        for piece in _chunk(page_text, chunk=chunk_size, overlap=overlap, encoder=tokenizer):
            chunks.append({"file": name, "page": page_number, "text": piece})

def parse_and_chunk(filename: str, data: bytes) -> List[Dict[str, object]]:
    """Parse ``data`` according to ``filename`` extension and chunk the text."""

    ext = filename.rsplit(".", 1)[-1].lower()
    if ext == "pdf":
        pages = list(_iter_pdf_text(data))
    elif ext == "docx":
        pages = list(_iter_docx_text(data))
    elif ext == "txt":
        pages = list(_iter_txt_text(data))
    else:
        return []

    if not pages:
        return []

    chunk_size = _normalise_window_size(int(os.getenv("RAG_CHUNK", "900")))
    overlap = _normalise_overlap(chunk_size, int(os.getenv("RAG_OVERLAP", "140")))
    tokenizer = _get_tokenizer()

    chunks: List[Dict[str, object]] = []
    for page, text in pages:
        for piece in _chunk(text, chunk=chunk_size, overlap=overlap, encoder=tokenizer):
            chunks.append({"file": filename, "page": page, "text": piece})
        main
    return chunks


__all__ = ["_chunk", "_clean", "_get_tokenizer", "parse_and_chunk"]
